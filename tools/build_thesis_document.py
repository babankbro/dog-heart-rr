from pathlib import Path
import csv, json, math, re, statistics

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "out" / "thesis_ekg_yolo_preprocessing.docx"
FONT = "TH Sarabun New"
BLUE = "1F4E79"
LIGHT = "D9EAF7"


def font(run, size=16, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_width(cell, dxa):
    tcw = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
    if tcw is None:
        tcw = OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tcw)
    tcw.set(qn("w:w"), str(dxa)); tcw.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblpr = table._tbl.tblPr
    tblw = tblpr.first_child_found_in("w:tblW")
    tblw.set(qn("w:w"), str(sum(widths))); tblw.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd"); ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa"); tblpr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(w)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells): set_cell_width(cell, widths[i])


def style_table(table, header=True):
    for ri, row in enumerate(table.rows):
        trpr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri == 0 and header: set_cell_shading(cell, LIGHT)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
                for r in p.runs: font(r, 12, bold=(ri == 0))
    if header:
        trpr = table.rows[0]._tr.get_or_add_trPr(); rep = OxmlElement("w:tblHeader"); rep.set(qn("w:val"), "true"); trpr.append(rep)


def add_table(doc, rows, widths=None):
    if not rows: return
    t = doc.add_table(rows=0, cols=len(rows[0])); t.style = "Table Grid"
    for vals in rows:
        cells = t.add_row().cells
        for i, v in enumerate(vals): cells[i].text = str(v)
    if widths is None:
        widths = [9360 // len(rows[0])] * len(rows[0])
        widths[-1] += 9360 - sum(widths)
    else:
        widths = list(widths)
    set_table_geometry(t, widths); style_table(t)
    doc.add_paragraph()


def new_decimal_num(doc):
    root = doc.part.numbering_part.element
    abs_ids = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    aid = max(abs_ids, default=0) + 1; nid = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum"); abstract.set(qn("w:abstractNumId"), str(aid))
    lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), "0")
    for tag, attr, val in [("w:start","w:val","1"),("w:numFmt","w:val","decimal"),("w:lvlText","w:val","%1."),("w:lvlJc","w:val","left")]:
        e=OxmlElement(tag); e.set(qn(attr),val); lvl.append(e)
    ppr=OxmlElement("w:pPr"); tabs=OxmlElement("w:tabs"); tab=OxmlElement("w:tab"); tab.set(qn("w:val"),"num"); tab.set(qn("w:pos"),"540"); tabs.append(tab)
    ind=OxmlElement("w:ind"); ind.set(qn("w:left"),"540"); ind.set(qn("w:hanging"),"280"); ppr.extend([tabs,ind]); lvl.append(ppr); abstract.append(lvl); root.append(abstract)
    num=OxmlElement("w:num"); num.set(qn("w:numId"),str(nid)); ref=OxmlElement("w:abstractNumId"); ref.set(qn("w:val"),str(aid)); num.append(ref); root.append(num)
    return nid


def add_numbered_block(doc, items):
    # Use explicit numerals so Word cannot silently continue numbering from a
    # previous list when it normalizes equivalent numbering definitions.
    for index, text in enumerate(items, start=1):
        p=doc.add_paragraph(f"{index}.  {clean_md(text)}")
        p.paragraph_format.left_indent=Inches(0.32)
        p.paragraph_format.first_line_indent=Inches(-0.22)
        p.paragraph_format.space_after=Pt(4)


def clean_md(s):
    s = "".join(ch for ch in s if ch in "\t\n\r" or ord(ch) >= 32)
    s = re.sub(r"\[\[([^]|]+)(?:\|([^]]+))?\]\]", lambda m: m.group(2) or m.group(1), s)
    s = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1", s)
    s = s.replace("**", "").replace("`", "").replace("$", "")
    return s


def add_markdown(doc, path, chapter_title=None):
    lines = path.read_text(encoding="utf-8").splitlines(); i = 0
    if chapter_title:
        doc.add_page_break(); doc.add_heading(chapter_title, level=1)
    while i < len(lines):
        line = lines[i].strip()
        if not line or line == "---": i += 1; continue
        if line.startswith("```mermaid"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"): i += 1
            i += 1
            continue
        if line.startswith("# "):
            if not chapter_title: doc.add_heading(clean_md(line[2:]), level=1)
            i += 1; continue
        if line.startswith("## "): doc.add_heading(clean_md(line[3:]), level=2); i += 1; continue
        if line.startswith("### "): doc.add_heading(clean_md(line[4:]), level=3); i += 1; continue
        if line.startswith("![["):
            name = line[3:-2]; img = ROOT / "vault" / "assets" / name
            if not img.exists(): img = ROOT / "out" / "figures" / name
            if img.exists():
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sizes = {"fig-pipeline.png": 3.0, "fig-crop-pre.png": 1.25, "fig-point-pre.png": 4.0, "fig-r-position.png": 5.8, "fig-domain-gap.png": 5.8}
                p.add_run().add_picture(str(img), width=Inches(sizes.get(name, 5.5)))
                cap = doc.add_paragraph(f"ภาพประกอบ: {name}"); cap.style = "Caption"; cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1; continue
        if line.startswith("|"):
            block=[]
            while i < len(lines) and lines[i].strip().startswith("|"):
                vals=[clean_md(x.strip()) for x in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"[-:]+", x) for x in vals): block.append(vals)
                i += 1
            if block: add_table(doc, block)
            continue
        if line.startswith("$$"):
            eq=[]; line=line[2:]
            while True:
                if "$$" in line: eq.append(line.split("$$")[0]); break
                eq.append(line)
                i += 1
                if i >= len(lines): break
                line=lines[i].strip()
            p=doc.add_paragraph(clean_md(" ".join(eq))); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: font(r, 14, italic=True)
            i += 1; continue
        if re.match(r"^\d+\. ", line):
            # Preserve the source numeral because list items may contain
            # continuation paragraphs between successive numbered entries.
            p=doc.add_paragraph(clean_md(line))
            p.paragraph_format.left_indent=Inches(0.32)
            p.paragraph_format.first_line_indent=Inches(-0.22)
            p.paragraph_format.space_after=Pt(4)
            i += 1
            continue
        if re.match(r"^[-*] ", line):
            text = re.sub(r"^[-*]\s+", "", line); p=doc.add_paragraph(clean_md(text), style="List Bullet")
        elif line.startswith(">"):
            p=doc.add_paragraph(clean_md(line.lstrip("> "))); p.style="Quote"
        elif line.startswith("```"):
            i += 1; code=[]
            while i < len(lines) and not lines[i].strip().startswith("```"): code.append(lines[i]); i += 1
            p=doc.add_paragraph("\n".join(code)); p.style="Quote"
        else:
            parts=[line]; j=i+1
            while j < len(lines) and lines[j].strip() and not re.match(r"^(#|\||!\[\[|```|\$\$|[-*] |\d+\. )", lines[j].strip()):
                parts.append(lines[j].strip()); j += 1
            p=doc.add_paragraph(clean_md(" ".join(parts))); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; i=j-1
        i += 1


def training_summary():
    rows=[["แบบจำลอง", "Epoch ที่ mAP50-95 สูงสุด", "Precision", "Recall", "mAP50", "mAP50-95"]]
    for label, path in [("Crop YOLO11m", ROOT/"artifacts/yolo11m_crop_clean-2/results.csv"), ("Point YOLO11m", ROOT/"artifacts/yolo11m_point_clean-3/results.csv")]:
        data=list(csv.DictReader(path.open(encoding="utf-8-sig")))
        key="metrics/mAP50-95(B)"; best=max(data, key=lambda x: float(x[key]) if x[key] not in ("nan","") else -1)
        rows.append([label,best["epoch"],f'{float(best["metrics/precision(B)"]):.3f}',f'{float(best["metrics/recall(B)"]):.3f}',f'{float(best["metrics/mAP50(B)"]):.3f}',f'{float(best[key]):.3f}'])
    return rows


def patient_rr_rows():
    idx=json.loads((ROOT/"data/patients.json").read_text(encoding="utf-8"))["patients"]
    meta={p["id"]:p for p in idx}; agg={pid:{"imgs":set(),"rr":[],"bpm":[],"r":0,"model":0,"anchor":0,"flags":0} for pid in meta}
    for f in (ROOT/"out/results").glob("*.json"):
        d=json.loads(f.read_text(encoding="utf-8")); pid=d.get("image","").split("/")[0]
        if pid not in agg: continue
        a=agg[pid]; a["imgs"].add(d["image"]); st=d.get("result",{}).get("stats",{})
        a["r"]+=int(st.get("n_peaks",0)); a["model"]+=int(st.get("n_model",0)); a["anchor"]+=int(st.get("n_anchor",0))
        for row in d.get("rows",[]):
            if row.get("rr_mm") not in ("",None): a["rr"].append(float(row["rr_mm"]))
            if row.get("bpm") not in ("",None): a["bpm"].append(float(row["bpm"]))
            if row.get("flag"): a["flags"]+=1
    rows=[["รหัส", "ชื่อ", "กลุ่ม", "ภาพ", "R", "RR n", "RR เฉลี่ย±SD (mm)", "BPM เฉลี่ย", "Model/Anchor", "Flag"]]
    for pid in meta:
        p=meta[pid]; a=agg[pid]; rr=a["rr"]; bpm=a["bpm"]
        rr_text="—" if not rr else f"{statistics.mean(rr):.2f} ± {(statistics.stdev(rr) if len(rr)>1 else 0):.2f}"
        bpm_text="—" if not bpm else f"{statistics.mean(bpm):.1f}"
        rows.append([pid,p.get("name") or "—",p.get("group") or "—",len(a["imgs"]),a["r"],len(rr),rr_text,bpm_text,f'{a["model"]}/{a["anchor"]}',a["flags"]])
    return rows


def main():
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(1); sec.right_margin=Inches(1)
    sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
    styles=doc.styles
    normal=styles["Normal"]; normal.font.name=FONT; normal.font.size=Pt(16); normal.paragraph_format.space_after=Pt(8); normal.paragraph_format.line_spacing=1.333
    for name,size,color,before,after in [("Heading 1",20,BLUE,18,10),("Heading 2",18,BLUE,12,6),("Heading 3",16,"1F4D78",8,4)]:
        s=styles[name]; s.font.name=FONT; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    for sname in ["Caption","Quote","List Bullet","List Number"]:
        styles[sname].font.name=FONT; styles[sname].font.size=Pt(14)
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(header.add_run("เอกสารร่างสำหรับวิทยานิพนธ์ | ระบบตรวจหา R peak จากภาพ EKG"),10,color="666666")
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(footer.add_run("งานวิจัยและพัฒนาระบบประมวลผลภาพคลื่นไฟฟ้าหัวใจ"),10,color="777777")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(100)
    font(p.add_run("การพัฒนาระบบตรวจหายอดคลื่นอาร์\nจากภาพคลื่นไฟฟ้าหัวใจด้วย YOLO11\nและการประมวลผลภาพแบบผสมผสาน"),24,bold=True,color=BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("เอกสารเนื้อหา ทฤษฎี วิธีดำเนินการ และร่างผลการทดลองสำหรับวิทยานิพนธ์"),16)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(80); font(p.add_run("จัดทำจากซอร์สโค้ด ผลการทดลอง และผลลัพธ์ที่บันทึกในโครงการ\nข้อมูล ณ วันที่ 2 กันยายน 2569"),14,color="555555")
    doc.add_page_break(); doc.add_heading("สารบัญเนื้อหา",1)
    for x in ["บทที่ 2 ทฤษฎีและงานที่เกี่ยวข้อง", "บทที่ 3 วิธีดำเนินการวิจัย", "บทที่ 4 ผลการทดลอง preprocessing", "ภาคผนวก ก ผล validation ของแบบจำลอง", "ภาคผนวก ข ตาราง RR รายตัว"]: doc.add_paragraph(x,style="List Bullet")
    add_markdown(doc, ROOT/"vault/09 - พื้นฐานและวิวัฒนาการของ YOLO.md", "บทที่ 2 ทฤษฎีและความรู้พื้นฐานเกี่ยวกับ YOLO")
    add_markdown(doc, ROOT/"vault/01 - Methodology.md", "บทที่ 3 วิธีดำเนินการวิจัย")
    add_markdown(doc, ROOT/"vault/10 - ผลการทดลอง Preprocessing.md", "บทที่ 4 ผลการทดลองการเตรียมภาพ")
    doc.add_page_break(); doc.add_heading("ภาคผนวก ก ผล validation ของแบบจำลอง",1)
    doc.add_paragraph("ตารางนี้สรุป epoch ที่ให้ mAP50-95 สูงสุดจากไฟล์ results.csv ของการฝึกแต่ละแบบจำลอง ค่าดังกล่าวเป็นผลบน validation set ของชุดฝึก ไม่ใช่ความแม่นยำของ preprocessing บนชุดภาพสุนัข 67 ภาพ")
    add_table(doc,training_summary(),[2200,1700,1300,1300,1300,1560])
    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape.orientation = 1
    landscape.page_width, landscape.page_height = sec.page_height, sec.page_width
    landscape.top_margin=Inches(.7); landscape.bottom_margin=Inches(.7); landscape.left_margin=Inches(1); landscape.right_margin=Inches(1)
    doc.add_heading("ภาคผนวก ข ตารางผล RR แยกรายตัว",1)
    doc.add_paragraph("คำนวณจากไฟล์ JSON ที่บันทึกอยู่ใน out/results จำนวน 67 ภาพ โดยรวมช่วง RR ที่มีค่าภายในทุกภาพของสัตว์แต่ละตัว ค่า BPM เป็นค่าเฉลี่ยของช่วงที่รายงาน ระบบยังไม่มี ground truth จากการวัดด้วยมือ จึงเป็นผลเชิงระบบที่ต้องผ่านการตรวจสอบก่อนตีความทางคลินิก")
    add_table(doc,patient_rr_rows(),[900,1000,2700,700,700,700,1900,1100,1600,1660])
    OUT.parent.mkdir(exist_ok=True); doc.save(OUT); print(OUT)

if __name__ == "__main__": main()
